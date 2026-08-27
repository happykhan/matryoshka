#!/usr/bin/env python3
"""Build the self-contained Tn1/Tn2/Tn3 evidence report.

The report deliberately embeds the accession-specific figures and tables. It
is intended to be readable without browsing the output directories.
"""

from __future__ import annotations

import csv
import json
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).parents[1]
OUTPUT = REPO / "reports" / "Matryoshka-Tn123-definition-and-validation-report.docx"
REVIEWED = REPO / "demo-output" / "tn123-reviewed-definitions" / "proof-bundle"
PEK499 = REPO / "demo-output" / "tn123-pEK499" / "proof-bundle"
ARBITRARY = REPO / "demo-output" / "arbitrary-tn123" / "proof-bundle"
LEDGER = REPO / "docs" / "validation" / "tn123-real-accession-results.tsv"

INK = "17212B"
NAVY = "17365D"
BLUE = "2E74B5"
PALE_BLUE = "EEF5FB"
PALE_GRAY = "F2F4F7"
MUTED = "5B6570"
GREEN = "DFF2E6"
GREEN_TEXT = "17633A"

ACCESSION_ORDER = [
    "Tn1_NC_008357",
    "Tn2_AY123253",
    "Tn3_HM749966",
    "Tn2c_HM749967",
    "Tn2_1_CP028717",
    "Tn1Mer_GQ160960",
    "Tn3_V00613",
]


def font(run, size: float = 10.5, *, bold: bool = False, color: str = INK,
         name: str = "Aptos", italic: bool = False) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top: int = 80, start: int = 120,
                 bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths: list[int], *, indent: int = 120) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for tag, value in (("w:tblW", total), ("w:tblInd", indent)):
        node = tbl_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths: list[int], *, font_size: float = 8.8,
              green_last: bool = False) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table_geometry(table, widths)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    header_properties.append(header_flag)
    for col, heading in enumerate(headers):
        cell = table.cell(0, col)
        shade(cell, PALE_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(heading)
        font(run, font_size, bold=True, color=NAVY)
    for row_index, values in enumerate(rows, start=1):
        for col, value in enumerate(values):
            cell = table.cell(row_index, col)
            if green_last and col == len(values) - 1 and value == "PASS":
                shade(cell, GREEN)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            font(run, font_size, bold=(green_last and value == "PASS"),
                 color=GREEN_TEXT if green_last and value == "PASS" else INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_para(doc: Document, text: str = "", *, size: float = 10.5,
             bold: bool = False, color: str = INK, after: float = 6,
             before: float = 0, align=None, italic: bool = False,
             keep_with_next: bool = False) -> object:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_with_next = keep_with_next
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    font(run, size, bold=bold, color=color, italic=italic)
    return p


def add_code(doc: Document, text: str, *, size: float = 8.0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    font(run, size, name="Aptos Mono", color=INK)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_BLUE)
    p._p.get_or_add_pPr().append(shd)


def add_bullet(doc: Document, text: str, *, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    font(run, 10.2)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    font(run, 8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_section(section, *, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.clear()
    run = header.add_run("MATRYOSHKA  |  Tn1/Tn2/Tn3 evidence report")
    font(run, 8.0, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.clear()
    set_page_number(footer)


def new_section(doc: Document, *, landscape: bool = False) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=landscape)


def configure_document(doc: Document) -> None:
    configure_section(doc.sections[0], landscape=False)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Title", 25, NAVY, 0, 8),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 17, BLUE, 15, 7),
        ("Heading 2", 13.5, NAVY, 11, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Bullet 2"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1


def set_alt_text(shape, title: str, description: str) -> None:
    doc_pr = shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_svg(doc: Document, svg: Path, temp_dir: Path, *, width: float,
            title: str, description: str, after: float = 2) -> None:
    png = temp_dir / f"{svg.parent.name}-{svg.stem}.png"
    subprocess.run(
        ["/opt/homebrew/bin/rsvg-convert", "-w", "2400", "-o", str(png), str(svg)],
        check=True,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    shape = p.add_run().add_picture(str(png), width=Inches(width))
    set_alt_text(shape, title, description)


def figure_label(doc: Document, text: str) -> None:
    add_para(doc, text, size=8.5, color=MUTED, after=3, italic=True,
             keep_with_next=True)


def load_proof(bundle: Path) -> dict:
    return json.loads((bundle / "proof" / "proof.json").read_text(encoding="utf-8"))


def proof_loci(proof: dict) -> dict[str, list[dict]]:
    return {record["id"]: record.get("loci", []) for record in proof["records"]}


def add_cover(doc: Document) -> None:
    add_para(doc, "TECHNICAL EVIDENCE REPORT", size=10, bold=True, color=BLUE,
             after=18, before=54, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Automated detection and visualisation\nof Tn1, Tn2 and Tn3")
    add_para(doc, "Expert rules, real-accession validation, locus maps,\nlocus tables and hierarchical outputs",
             size=13, color=MUTED, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_table(
        doc,
        ["Definition set", "Software status", "Report date", "Validation"],
        [["2026-08-27", "External alpha", "27 August 2026", "14/14 checks pass"]],
        [2100, 2100, 2100, 2100],
        font_size=9.5,
    )
    add_para(doc, "Purpose", size=12, bold=True, color=NAVY, before=16, after=5)
    add_para(doc, "This document is the reviewable evidence package: it explains what is run, how the biological rules are specified, and what the program produced from the reviewed accession sequences. The actual figures and tables are embedded in the report.", size=11, after=10)
    add_para(doc, "Scope statement", size=12, bold=True, color=NAVY, before=10, after=5)
    add_para(doc, "The validated scope is Tn1, Tn2, Tn3 and the reviewed close definitions listed here. It is not yet an exhaustive replacement for expert mobile-element annotation across every family.", size=11)


def add_contents(doc: Document) -> None:
    new_section(doc)
    add_heading(doc, "Contents and claims at a glance", 1)
    rows = [
        ["1", "What the tool does", "Input, analysis stages, calls and outputs"],
        ["2", "What is actually run", "Built-in BLAST scans; optional ISEScan, AMRFinderPlus and IntegronFinder"],
        ["3", "Expert-rule architecture", "Human-readable YAML, schema validation and extension workflow"],
        ["4", "Worked Tn1 definition", "Grammar, thresholds, exact naming and variants"],
        ["5", "Validation design", "Real positive controls, natural fragments, variants and negative controls"],
        ["6", "Real-accession evidence", "locus map, locus table, hierarchy and component ledger for every reviewed definition"],
        ["7", "Arbitrary-sequence demonstration", "Three distinct loci recovered from one 44.6 kb contig"],
        ["8", "Limitations and next work", "What is proven, what remains expert review, and what is not claimed"],
    ]
    add_table(doc, ["Section", "Question answered", "Evidence"], rows,
              [900, 2700, 4860], font_size=9.2)
    add_heading(doc, "What this report demonstrates", 2)
    for text in (
        "An arbitrary multi-FASTA input can be scanned without pre-existing annotations.",
        "Tn1/Tn2/Tn3 components are detected independently and assembled using an explicit component grammar.",
        "A complete locus is classified from its component grammar and weighted local component profiles; a whole-element comparison is secondary confirmation.",
        "Every result can be inspected as a locus map, a locus table, the original hierarchy view, JSON, GFF3 and CellGen format.",
        "The seven reviewed accession definitions produce complete component grammars and exact declared calls; pEK499 produces two retained partial Tn2-like fragments rather than an invented complete Tn2.",
    ):
        add_bullet(doc, text)
    add_heading(doc, "Important interpretation", 2)
    add_para(doc, "The seven exact accession runs are positive controls against the sequences used to define those names. They prove correct rule execution, component recovery and output generation. They do not by themselves estimate sensitivity on a broad independent population. The natural pEK499 fragments, synthetic variation tests and related-element negative controls test behaviour away from exact references.")


def add_tool_explanation(doc: Document) -> None:
    new_section(doc)
    add_heading(doc, "1. What the tool does", 1)
    add_para(doc, "Matryoshka is a per-sequence nested mobile-element annotator. For this report, its validated task is to find Tn1/Tn2/Tn3-group components in a FASTA sequence, assemble candidate loci, assign a reviewed name only when the evidence satisfies the expert rules, and render the same evidence in several complementary formats.")
    add_heading(doc, "Input to output", 2)
    steps = [
        ["1", "Read FASTA", "One or many contigs or complete sequences; no feature annotation is required."],
        ["2", "Detect components", "Scan independently for 38 bp terminal inverted repeats, blaTEM, tnpR, res and tnpA."],
        ["3", "Assemble loci", "Group nearby features and test the required order and orientation on either strand."],
        ["4", "Classify by expert rules", "Score the detected component profiles and assign a type only when the declared threshold and margin are met."],
        ["5", "Compare references", "Optionally confirm an exact reviewed definition and record secondary closest-reference context."],
        ["6", "Write outputs", "JSON, GFF3, CellGen, hierarchy SVG, locus map, locus table and proof ledger."],
    ]
    add_table(doc, ["Stage", "Operation", "Why it matters"], steps,
              [700, 2100, 5660], font_size=9.2)
    add_heading(doc, "Call classes", 2)
    add_table(doc, ["Visible result", "Meaning"], [
        ["Tn1 / Tn2 / Tn3 or named subtype", "Complete component grammar and exact match to a reviewed declared definition."],
        ["Tn1-like / Tn2-like / Tn3-like", "Complete grammar and a clear best weighted component profile, without requiring a complete-element match."],
        ["Tn1/2/3 fragment", "At least 400 aligned reference bases, but an end or required component is missing or the grammar is incomplete."],
        ["Unresolved Tn1/Tn2/Tn3-group unit", "The group grammar is supported but the component-profile score or type margin is insufficient."],
        ["No Tn1/Tn2/Tn3 call", "Some shared components may still be annotated and drawn, but the element-level rule is not satisfied."],
    ], [2800, 5660], font_size=9.2)


def add_run_details(doc: Document) -> None:
    new_section(doc)
    add_heading(doc, "2. What is actually run", 1)
    add_heading(doc, "Built-in analysis", 2)
    add_para(doc, "The command scans independently for the component sequences and assembles their collinear matches across substitutions, insertions and deletions. The expert-rule classifier uses those components to establish the family and type. Under the default validated profile, a separate whole-element comparison is then attached as confirmation or context. The tn123-components profile omits that complete-element lookup entirely.")
    add_heading(doc, "Optional detector layer", 2)
    add_para(doc, "The --detectors option controls additional programs, not the built-in Tn1/Tn2/Tn3 scan. The report bundles were generated with --detectors none, which means ISEScan, AMRFinderPlus and IntegronFinder were not launched. Built-in BLAST detection and the expert-rule classifier still ran. With --detectors available or --detectors all, those programs add broader IS, AMR-gene and integron evidence.")
    add_table(doc, ["Analysis", "Role", "Used for this report"], [
        ["BLAST component scan", "Detect IRL/IRR, blaTEM, tnpR, res and tnpA independently", "Yes; primary"],
        ["Component assembler", "Test count, order, orientation and ends", "Yes"],
        ["Expert-rule classifier", "Assign family/type-like, fragment or unresolved result from weighted component profiles", "Yes; primary"],
        ["BLAST whole-locus comparison", "Optionally confirm an exact reviewed definition or add closest-reference context", "Yes; secondary"],
        ["Boundary/direct-repeat check", "Record evidence immediately outside inferred ends when flanks exist", "Yes"],
        ["ISEScan / AMRFinderPlus / IntegronFinder", "Broader optional component discovery", "Not launched"],
    ], [2600, 4100, 1760], font_size=9.1)
    add_heading(doc, "Reproducible commands", 2)
    add_code(doc, "# Prove discovery without complete-element lookup\nmatryoshka run arbitrary-demo.fasta --out results/component-only \\\n  --profile tn123-components --detectors none\n\n# Seven reviewed definitions with secondary exact confirmation\nmatryoshka run matryoshka/references/tn1_tn2_tn3.fasta \\\n  --out results/tn123-reviewed --detectors none\n\n# Export the rules for expert review\nmatryoshka definitions --format markdown --out tn123-definitions-review.md")
    add_para(doc, "Reproducibility metadata", size=11.5, bold=True, color=NAVY, before=8, after=4)
    add_para(doc, "Each result directory includes run.json, the software and definition versions, reference profile, detector selection, input checksum, runtime parameters and a direct index of every locus map/table pair.")


def add_rules(doc: Document) -> None:
    new_section(doc)
    add_heading(doc, "3. How the expert rules are specified", 1)
    add_para(doc, "The biological definitions are data, not a collection of element-specific if-statements hidden in the Python source. The authoritative file is matryoshka/tn123_definitions.yaml. A JSON Schema checks its required structure, and the same content can be exported as YAML, JSON or a formatted Markdown review document.")
    add_heading(doc, "Definition structure", 2)
    add_table(doc, ["Block", "Human meaning"], [
        ["family", "Shared biological description, expected terminal-IR length and target-site duplication length."],
        ["grammar", "Required component counts and forward/reverse order."],
        ["component_detection", "Identity, coverage, length and chaining thresholds for each component class."],
        ["classification", "Weighted component-profile type rules, exact-reference confirmation and fragment rules."],
        ["types", "Canonical Tn1, Tn2 and Tn3 sequences and their annotated component layouts."],
        ["definitions", "Reviewable named sequences/subtypes, accession provenance, expert description and differences from parent."],
        ["related_element_policy", "Explicit negative-control policy to prevent shared components being over-named."],
    ], [2400, 6060], font_size=9.2)
    add_heading(doc, "Shared Tn1/Tn2/Tn3 grammar", 2)
    add_code(doc, "terminal IR -> blaTEM -> tnpR -> res -> tnpA -> terminal IR\n    38 bp       one       one      one     one        38 bp")
    add_para(doc, "The reverse-complement copy must contain the corresponding reversed order. This establishes group membership, but it does not distinguish Tn1, Tn2 and Tn3 because their structures and much of their sequence are shared.")
    add_heading(doc, "Current thresholds", 2)
    add_table(doc, ["Decision", "Rule"], [
        ["Exact declared definition", "100% identity, 100% reference coverage, both ends, zero mismatches/insertions/deletions and complete component grammar."],
        ["Rule-based type", "Complete grammar; weighted local component score at least 90%; best type exceeds the next by at least 0.5 points."],
        ["Component weights", "blaTEM 1; tnpR 2; res 3; tnpA 3. The combined profile, rather than a single gene, assigns the type-like call."],
        ["Secondary reference", "At least 95% identity and 80% coverage; used for context or exact confirmation, never to create the family/type."],
        ["Fragment", "At least 400 aligned reference bases; retained without a complete name."],
    ], [2600, 5860], font_size=9.2)


def add_tn1_example(doc: Document) -> None:
    new_section(doc)
    add_heading(doc, "4. Worked expert definition: Tn1", 1)
    add_para(doc, "Human-readable rule", size=11.5, bold=True, color=NAVY, after=4)
    add_para(doc, "Call a Tn1-like candidate when the six required components are independently detected in the correct order and their weighted local profiles best support Tn1 above the declared score and margin. Upgrade that call to exact reviewed Tn1 only when an optional secondary comparison is a complete, gap-free and mismatch-free match to NC_008357. The canonical sequence carries blaTEM-2.", size=11)
    add_heading(doc, "YAML representation used by the program", 2)
    add_code(doc, "classification:\n  type_assignment:\n    method: weighted component profiles\n    discriminator_role_weights:\n      blaTEM: 1\n      tnpR: 2\n      res: 3\n      tnpA: 3\n    minimum_component_profile_score_percent: 90\n    minimum_type_margin_percent: 0.5\n  reference_comparison:\n    role: secondary context after rule-based classification\n\ntypes:\n  Tn1:\n    canonical_reference: Tn1_NC_008357\n    source_accession: NC_008357\n    components:\n      - {role: terminal_IR, name: IRL, start: 1, end: 38}\n      - {role: blaTEM, name: blaTEM-2, start: 148, end: 1008, strand: '-'}\n      - {role: tnpR, name: tnpR, start: 1191, end: 1748, strand: '-'}\n      - {role: res, name: res, start: 1754, end: 1867}\n      - {role: tnpA, name: tnpA, start: 1911, end: -34, strand: '+'}\n      - {role: terminal_IR, name: IRR, start: -38, end: -1}", size=7.3)
    add_para(doc, "Negative coordinates count from the right-hand end of the reference. This keeps the component layout readable and reusable for definitions of different lengths.", size=9.5, color=MUTED)
    add_heading(doc, "How to add or revise a subtype", 2)
    for text in (
        "Add the public reference sequence and provenance.",
        "Add a definitions entry containing the parent type, visible name, accession, review status, expert rule and known differences.",
        "Declare any reviewed inserted or nested context under additional_features.",
        "Rebuild the reference FASTA and checksum manifest.",
        "Add a positive accession test and at least one related-element negative control.",
        "Rerun the validation ledger and complete output-bundle test.",
    ):
        add_bullet(doc, text)


def add_boundaries_and_validation(doc: Document) -> None:
    new_section(doc)
    add_heading(doc, "5. Boundaries, direct repeats and validation design", 1)
    add_heading(doc, "How direct repeats are detected", 2)
    add_para(doc, "For this group, a five-base direct repeat is supporting boundary evidence, not the naming rule. After candidate element ends have been inferred, Matryoshka reads the five bases immediately outside the left end and the five bases immediately outside the right end. An identical pair is reported with sequence and coordinates. If the record ends at the element boundary, the test is marked untestable because the flank is absent. If both flanks exist but differ, the result records that a repeat was searched for and not found.")
    add_para(doc, "A five-base repeat can occur by chance. It therefore cannot create a Tn1/Tn2/Tn3 name without the complete component grammar and type-profile evidence.", size=9.7, color=MUTED)
    add_heading(doc, "Validation set", 2)
    add_table(doc, ["Evidence class", "Examples", "Purpose"], [
        ["Reviewed exact definitions", "NC_008357, AY123253, HM749966, HM749967, CP028717, GQ160960, V00613", "Positive controls for type/subtype execution and output generation."],
        ["Natural partial sequence", "pEK499, EU935739.1", "Verify that two separated Tn2-related fragments remain partial."],
        ["Minor sequence variation", "Tn1 with 25 substitutions", "Verify a complete non-identical sequence becomes Tn1-like."],
        ["Structural variation", "Tn2 with an 800 bp insertion", "Verify an inserted sequence is retained and the call remains qualified."],
        ["Internal deletion", "Tn2 with a 300 bp deletion", "Verify split component evidence is chained and remains Tn2-like."],
        ["Truncation", "Tn1 missing 700 left-hand bases", "Verify the sequence becomes a fragment rather than exact Tn1."],
        ["Related elements", "Tn1696, Tn1721, Tn5403", "Verify shared components do not force a Tn1/Tn2/Tn3 name."],
    ], [2400, 2700, 3360], font_size=8.8)
    add_heading(doc, "Headline result", 2)
    add_para(doc, "All 14 real-accession ledger checks and all 23 focused Tn1/Tn2/Tn3 acceptance tests pass. The component-only tests type substituted, inserted and deleted variants without a complete-element lookup. The reviewed inputs receive exact confirmation; the pEK499 loci remain partial; related elements are not misnamed.", size=11, bold=True, color=GREEN_TEXT)


def add_validation_ledger(doc: Document) -> None:
    new_section(doc, landscape=True)
    add_heading(doc, "Validation ledger", 1)
    with LEDGER.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle, delimiter="\t"))
    compact = []
    for item in rows:
        compact.append([
            item.get("scenario", ""), item.get("record", ""), item.get("accession", ""),
            item.get("expected", ""), item.get("observed_call", ""), item.get("result", ""),
        ])
    add_table(doc, ["Scenario", "Record", "Accession", "Expected", "Observed", "Result"], compact,
              [2200, 2100, 1500, 3500, 2400, 700], font_size=7.5, green_last=True)
    add_para(doc, "The machine-readable source is docs/validation/tn123-real-accession-results.tsv. Blank accession/type fields in the related-element rows are intentional: those tests assert that no Tn1/Tn2/Tn3 name is assigned.", size=8.7, color=MUTED)


def accession_title(record_id: str, locus: dict) -> str:
    del record_id
    accession = locus["known_element_match"]["source_accession"]
    return f"Real accession result — {locus['call']} ({accession})"


def add_accession_results(doc: Document, temp_dir: Path) -> None:
    proof = load_proof(REVIEWED)
    by_record = proof_loci(proof)
    for number, record_id in enumerate(ACCESSION_ORDER, start=1):
        locus = by_record[record_id][0]
        match = locus["known_element_match"]
        assembly = locus["assembly"]
        new_section(doc, landscape=True)
        add_heading(doc, f"6.{number} {accession_title(record_id, locus)}", 1)
        add_table(doc, ["Input record", "Observed call", "Coordinates", "Whole-locus match", "Grammar", "Proof"], [[
            record_id,
            locus["call"],
            f"{locus['start']:,}–{locus['end']:,} ({locus['strand']})",
            f"{match['whole_locus_identity']:.1f}% identity; {match['whole_locus_coverage']:.1f}% coverage",
            f"{assembly['detected_component_count']}/6 detected; order valid",
            locus["verdict"],
        ]], [2300, 1500, 2100, 3000, 2600, 800], font_size=8.4, green_last=True)
        add_para(doc, match["expert_rule"], size=9.0, color=MUTED, after=4)
        figure_label(doc, f"locus map generated from the detected and assembled features for {record_id}.")
        add_svg(doc, REVIEWED / locus["outputs"]["locus_map"], temp_dir, width=9.25,
                title=f"{locus['call']} locus maps",
                description=f"locus map for accession {match['source_accession']} with detected components and legend.")
        if record_id == "Tn2_1_CP028717":
            new_section(doc, landscape=True)
            add_heading(doc, "6.5 Continued — Tn2.1 locus table", 1)
        figure_label(doc, "locus table generated from the same call and evidence ledger.")
        add_svg(doc, REVIEWED / locus["outputs"]["locus_table"], temp_dir, width=9.25,
                title=f"{locus['call']} locus table",
                description=f"locus feature table for accession {match['source_accession']}.")

        new_section(doc, landscape=True)
        add_heading(doc, f"6.{number} Evidence detail — {locus['call']} ({match['source_accession']})", 1)
        figure_label(doc, "Original scale-accurate hierarchical view. Parent-child nesting is preserved.")
        add_svg(doc, REVIEWED / locus["outputs"]["hierarchy"], temp_dir, width=9.25,
                title=f"{locus['call']} hierarchy",
                description=f"Original hierarchical annotation view for accession {match['source_accession']}.")
        components = [[
            item["role"], item["name"], f"{item['start']:,}–{item['end']:,}", item["strand"],
            f"{item['identity']:.2f}%", f"{item['coverage']:.1f}%", item["evidence_class"],
        ] for item in locus["components"]]
        add_table(doc, ["Role", "Feature", "Coordinates", "Strand", "Identity", "Coverage", "Evidence"],
                  components, [1500, 1800, 1900, 800, 1200, 1200, 3000], font_size=7.8)
        add_para(doc, f"Definition: {match['definition_id']} (version {match['definition_version']}). Variant status: {match['variant_status']}. Reference-projected required components: {assembly['reference_projected_component_count']}. Context features from the reviewed definition: {assembly['reference_projected_context_feature_count']}.", size=8.7, color=MUTED)
        if record_id == "Tn2_1_CP028717":
            nested_stem = "Tn2_1_CP028717__ISEcp1-associated_insertion__130-4153.svg"
            new_section(doc, landscape=True)
            add_heading(doc, "6.5 Nested locus — ISEcp1-associated insertion", 1)
            add_para(doc, "The interrupted Tn2.1 definition contains a separately indexed ISEcp1-associated locus. It is drawn independently as well as inside the whole Tn2.1 view.")
            add_svg(doc, REVIEWED / "locus-map" / nested_stem, temp_dir, width=9.25,
                    title="Tn2.1 nested ISEcp1-associated insertion",
                    description="locus map of the separately indexed ISEcp1-associated insertion within Tn2.1.")
            add_svg(doc, REVIEWED / "locus-table" / nested_stem, temp_dir, width=9.25,
                    title="Tn2.1 nested ISEcp1-associated insertion table",
                    description="locus feature table for the separately indexed ISEcp1-associated insertion within Tn2.1.")


def add_partial_results(doc: Document, temp_dir: Path) -> None:
    proof = load_proof(PEK499)
    loci = proof["records"][0]["loci"]
    new_section(doc, landscape=True)
    add_heading(doc, "7. Natural partial-fragment result — pEK499 (EU935739.1)", 1)
    rows = []
    for locus in loci:
        match = locus["known_element_match"]
        present = ", ".join(role for role, ok in locus["assembly"]["requirements"].items() if ok)
        rows.append([
            f"{locus['start']:,}–{locus['end']:,}", locus["call"], match["best_match"],
            f"{match['whole_locus_coverage']:.2f}%", present or "none", locus["verdict"],
        ])
    add_table(doc, ["Coordinates", "Observed call", "Closest type", "Reference coverage", "Detected required roles", "Proof"],
              rows, [1900, 2500, 1500, 1700, 3900, 900], font_size=8.4)
    add_para(doc, "This is the key natural-context behaviour: two separated Tn2-related regions are retained independently. Neither contains both terminal ends and all required components, so the software does not join them or call an exact Tn2.", size=10.2, bold=True, color=NAVY)
    hierarchy = PEK499 / loci[0]["outputs"]["hierarchy"]
    figure_label(doc, "Original hierarchy view for the complete 117,536 bp pEK499 record.")
    add_svg(doc, hierarchy, temp_dir, width=9.25, title="pEK499 hierarchy",
            description="Original hierarchy view of pEK499 showing all detected validated-profile features.")
    for index, locus in enumerate(loci, start=1):
        match = locus["known_element_match"]
        new_section(doc, landscape=True)
        add_heading(doc, f"7.{index} pEK499 fragment {locus['start']:,}–{locus['end']:,}", 1)
        add_para(doc, f"Observed call: {locus['call']}; closest reviewed type: {match['best_match']}; identity {match['whole_locus_identity']:.1f}%; reference coverage {match['whole_locus_coverage']:.2f}%; proof verdict {locus['verdict']}.", size=10.0)
        figure_label(doc, "locus map. Missing element ends/components remain visibly absent.")
        add_svg(doc, PEK499 / locus["outputs"]["locus_map"], temp_dir, width=9.25,
                title="pEK499 partial Tn1/Tn2/Tn3 locus map",
                description=f"locus partial-fragment map at pEK499 coordinates {locus['start']} to {locus['end']}.")
        figure_label(doc, "locus table for the partial call. The partial status and incomplete grammar are explicit.")
        add_svg(doc, PEK499 / locus["outputs"]["locus_table"], temp_dir, width=9.25,
                title="pEK499 partial Tn1/Tn2/Tn3 table",
                description=f"locus table for the pEK499 fragment at {locus['start']} to {locus['end']}.")


def add_arbitrary_demo(doc: Document, temp_dir: Path) -> None:
    proof = load_proof(ARBITRARY)
    record = proof["records"][0]
    new_section(doc, landscape=True)
    add_heading(doc, "8. Arbitrary-sequence demonstration", 1)
    add_para(doc, "A single 44,647 bp contig was constructed with unrelated flanking sequence and three loci: a Tn1-derived sequence with 25 substitutions, a Tn2-derived sequence with an approximately 800 bp internal insertion, and a reverse-complement Tn3-derived sequence. It was processed with --profile tn123-components, which excludes complete Tn1/Tn2/Tn3 reference lookup. Every biological call below therefore comes from detected components and expert rules.")
    rows = []
    for locus in record["loci"]:
        match = locus["known_element_match"]
        scores = locus.get("classification", {}).get("component_type_scores", {})
        if not scores:
            scores = locus.get("rule_evidence", {}).get("component_type_scores", {})
        type_call = match.get("rule_based_type_call") or locus.get("family")
        rows.append([
            f"{locus['start']:,}–{locus['end']:,}", locus["call"], locus["strand"],
            str(type_call), str(scores.get(str(type_call), "see JSON")),
            "none", locus["verdict"],
        ])
    add_table(doc, ["Coordinates", "Call", "Strand", "Rule type", "Component score", "Whole-element lookup", "Proof"],
              rows, [2100, 1900, 800, 1200, 1300, 1300, 900], font_size=8.5, green_last=True)
    figure_label(doc, "Original hierarchy view: all three loci are retained on their shared contig and nested features remain visible.")
    add_svg(doc, ARBITRARY / record["loci"][0]["outputs"]["hierarchy"], temp_dir, width=9.25,
            title="Arbitrary contig hierarchy",
            description="Original hierarchical view of a 44.6 kb arbitrary contig containing three Tn1/Tn2/Tn3 loci.")
    for index, locus in enumerate(record["loci"], start=1):
        new_section(doc, landscape=True)
        add_heading(doc, f"8.{index} Arbitrary contig locus — {locus['call']}", 1)
        match = locus["known_element_match"]
        add_para(doc, f"The component grammar is complete and the expert component-profile rules assign {match.get('rule_based_type_call')}. No complete Tn1/Tn2/Tn3 reference comparison was run; the visible call is {locus['call']}. Insertions and other structural differences are retained from split component alignments and shown in the map/table.")
        figure_width = 8.65 if locus["call"] == "Tn2-like" else 9.25
        add_svg(doc, ARBITRARY / locus["outputs"]["locus_map"], temp_dir, width=figure_width,
                title=f"Arbitrary contig {locus['call']} locus map",
                description=f"locus map of the detected {locus['call']} locus on the arbitrary contig.")
        add_svg(doc, ARBITRARY / locus["outputs"]["locus_table"], temp_dir, width=figure_width,
                title=f"Arbitrary contig {locus['call']} locus table",
                description=f"locus evidence table for the detected {locus['call']} locus.")


def add_outputs_and_limits(doc: Document) -> None:
    new_section(doc)
    add_heading(doc, "9. Output package and interpretation", 1)
    add_table(doc, ["Output", "Purpose"], [
        ["annotation.json", "Complete machine-readable hierarchy, identifiers, coordinates, evidence, definitions and confidence."],
        ["annotation.gff3", "Interoperable genomic feature annotation."],
        ["annotation.cell", "Original nested CellGen/Wolvercote representation."],
        ["hierarchy/*.svg", "Original scale-accurate parent-child view of the record."],
        ["locus-map/*.svg", "Readable locus-centred locus map with its own legend/key."],
        ["locus-table/*.svg", "locus feature/evidence table for the same locus."],
        ["proof/report.html", "Linked human-readable evidence report."],
        ["proof/proof.json; components.tsv; matches.tsv", "Machine-checkable component-to-call ledger."],
        ["run.json", "Run metadata plus direct paths to every generated locus map and table."],
    ], [2700, 5760], font_size=9.0)
    add_heading(doc, "CellGen example", 2)
    add_code(doc, "({{}IRL, {}blaTEM-2, {}tnpR, {}res, {}tnpA, {}IRR}Tn1[family=\"Tn1\"])Tn1_NC_008357", size=8.3)
    add_para(doc, "The CellGen line carries the same nesting as the hierarchy graphic: the six components are children of Tn1, which belongs to the input record.")
    add_heading(doc, "How to read the figures", 2)
    for text in (
        "The locus maps prioritises biological readability around one locus.",
        "The locus table is the feature-by-feature description, including evidence and call qualification.",
        "The hierarchy view preserves scale and parent-child nesting across the whole record.",
        "The proof ledger separates sequence-detected required components from context added from a reviewed subtype definition.",
        "The legend/key is embedded in every locus map and table so the symbols are interpretable without prior experience of the notation.",
    ):
        add_bullet(doc, text)

    add_heading(doc, "10. Limitations and what should happen next", 1)
    add_para(doc, "What is proven now", size=11.5, bold=True, color=GREEN_TEXT, after=4)
    add_para(doc, "The implementation can recover the defined components, apply reviewable Tn1/Tn2/Tn3 rules, distinguish exact from qualified and partial calls, and generate the complete set of requested outputs on the included controls and demonstrations.")
    add_para(doc, "What is not yet proven", size=11.5, bold=True, color=NAVY, before=8, after=4)
    for text in (
        "Population-level sensitivity and specificity across a large independent accession panel.",
        "Exhaustive subtype coverage or complete parity with every component and compound element described across the source literature.",
        "Assigning a new formal element name to a previously undescribed structure without expert review; novel candidates within the defined family are deliberately reported as type-like.",
        "Target-site duplication evidence when the input sequence ends exactly at the element boundary or lacks sufficient flank.",
        "Consistent optional-detector availability on every installation; the built-in validated scan is the portable core.",
    ):
        add_bullet(doc, text)
    add_para(doc, "Recommended next validation", size=11.5, bold=True, color=NAVY, before=8, after=4)
    add_para(doc, "Have a domain expert review the human-readable definitions and embedded figures; freeze accepted naming decisions; then evaluate a blinded accession set containing true Tn1/Tn2/Tn3 examples, close variants, fragmented assemblies and related Tn3-family elements. Report locus-level precision/recall, boundary error and exact-versus-qualified naming accuracy.")
    add_heading(doc, "Expert decisions still requested", 2)
    for text in (
        "Confirm whether HM749967 should be displayed as Tn2c.",
        "Confirm the preferred name and internal labels for the CP028717 interrupted Tn2.1 structure.",
        "Confirm whether Tn1Mer is a public named subtype or an interrupted Tn1 example.",
        "Confirm the preferred wording for V00613 relative to canonical HM749966.",
        "Supply or approve additional type/subtype definitions and negative controls for the next expansion.",
    ):
        add_bullet(doc, text)


def build() -> None:
    for required in (REVIEWED, PEK499, ARBITRARY, LEDGER):
        if not required.exists():
            raise FileNotFoundError(required)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    add_tool_explanation(doc)
    add_run_details(doc)
    add_rules(doc)
    add_tn1_example(doc)
    add_boundaries_and_validation(doc)
    add_validation_ledger(doc)
    with tempfile.TemporaryDirectory(prefix="matryoshka-report-") as temp:
        temp_dir = Path(temp)
        add_accession_results(doc, temp_dir)
        add_partial_results(doc, temp_dir)
        add_arbitrary_demo(doc, temp_dir)
    add_outputs_and_limits(doc)
    doc.core_properties.title = "Automated detection and visualisation of Tn1, Tn2 and Tn3"
    doc.core_properties.subject = "Expert-rule architecture and accession-level evidence report"
    doc.core_properties.author = "Matryoshka project"
    doc.core_properties.keywords = "Tn1, Tn2, Tn3, locus maps, mobile genetic elements, validation"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
